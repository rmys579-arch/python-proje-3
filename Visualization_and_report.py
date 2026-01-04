import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt

#Visual settings (Readable font sizes)
sns.set_context("talk") # Makes text larger and more readable
plt.rcParams['axes.titlesize'] = 16

def add_spacer(doc):
    doc.add_paragraph("\n")

def create_report():
    df = pd.read_csv("cleaned_dataset.csv")
    doc = Document()
    
   # Report Title
    title = doc.add_heading('DATA ANALYSIS EXECUTIVE REPORT', 0)
    title.alignment = 1 # Center alignment
    
    doc.add_paragraph(f"This report contains an automated analysis of {len(df)} records. "
                      "Visualizations are powered by Seaborn for scientific accuracy.")

    num_cols = df.select_dtypes(include=['number']).columns.tolist()
    
    if num_cols:
        target = num_cols[0]
        
        # --- SECTION 1: DISTRIBUTION ---
        doc.add_heading('1. Data Distribution Analysis', level=1)
        plt.figure(figsize=(10, 5))
        sns.violinplot(data=df, x=target, color="lightgreen")
        plt.title(f"Detailed Distribution of {target}")
        plt.savefig("plot1.png", bbox_inches='tight')
        plt.close()
        
        doc.add_picture("plot1.png", width=Inches(5.5))
        doc.add_paragraph(f"The analysis of {target} shows the spread and density of the data. "
                          "The wider sections of the violin represent a higher probability of data points.")
        doc.add_page_break() # Her grafikten sonra yeni sayfaya geç

        # --- SECTION 2: CORRELATIONS (ADD TABLE) ---
        doc.add_heading('2. Correlation Insights', level=1)
        if len(num_cols) > 1:
            plt.figure(figsize=(8, 6))
            sns.heatmap(df.corr(numeric_only=True).round(2), annot=True, cmap='coolwarm')
            plt.savefig("plot2.png", bbox_inches='tight')
            plt.close()
            doc.add_picture("plot2.png", width=Inches(5.0))
            doc.add_paragraph("High correlation values (closer to 1 or -1) indicate a strong relationship.")

    doc.save("Final_Project_Report.docx")
    print("✨ SUCCESS: Professional report 'Final_Project_Report.docx' created!")

create_report()
